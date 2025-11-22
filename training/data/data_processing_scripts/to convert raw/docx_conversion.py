import os
from pathlib import Path
import docx
from docx.shared import RGBColor
import html

def extract_text_with_formatting_from_docx():
    """
    Extract text from DOCX files while preserving formatting.
    Converts to HTML format to retain bold, italics, highlighting and font sizes.
    """
    raw_path = Path("../raw")
    doc_output_path = Path("../processed/scripts")
    doc_output_path.mkdir(exist_ok=True, parents=True)
    
    doc_paths = list(raw_path.glob("**/*.docx"))
    
    print(f"Found {len(doc_paths)} DOCX files to process")
    
    for doc_path in doc_paths:
        txt_output_name = doc_path.stem + ".txt"
        html_output_name = doc_path.stem + ".html"
        txt_output_path = doc_output_path / txt_output_name
        html_output_path = doc_output_path / html_output_name
        
        if txt_output_path.exists() and html_output_path.exists():
            print(f"Files already extracted for {doc_path.name}, skipping...")
            continue
        
        print(f"Extracting text with formatting from {doc_path.name}...")
        
        try:
            doc = docx.Document(doc_path)
            
            html_content = "<html><head><style>"\
                          "mark {background-color: yellow;}\n"\
                          ".small {font-size: 0.8em;}\n"\
                          ".large {font-size: 1.2em;}\n"\
                          ".xlarge {font-size: 1.5em;}\n"\
                          ".xxlarge {font-size: 2em;}\n"\
                          "</style></head><body>"
            plain_text = ""
            
            for para in doc.paragraphs:
                if not para.text.strip():
                    html_content += "<br/>"
                    plain_text += "\n"
                    continue
                
                html_para = "<p>"
                para_text = ""
                
                for run in para.runs:
                    text = html.escape(run.text)
                    run_html = text
                    
                    if run.bold:
                        run_html = f"<strong>{run_html}</strong>"
                    if run.italic:
                        run_html = f"<em>{run_html}</em>"
                    if run.underline:
                        run_html = f"<u>{run_html}</u>"
                    
                    if run.font.highlight_color:
                        run_html = f"<mark>{run_html}</mark>"
                    
                    if hasattr(run.font, 'size') and run.font.size:
                        try:
                            size_pt = run.font.size.pt
                            if size_pt < 10:
                                run_html = f"<span class='small'>{run_html}</span>"
                            elif size_pt > 14 and size_pt <= 16:
                                run_html = f"<span class='large'>{run_html}</span>"
                            elif size_pt > 16 and size_pt <= 20:
                                run_html = f"<span class='xlarge'>{run_html}</span>"
                            elif size_pt > 20:
                                run_html = f"<span class='xxlarge'>{run_html}</span>"
                        except:
                            pass
                    
                    if run.font.color and run.font.color.rgb:
                        rgb = run.font.color.rgb
                        color_hex = f"#{rgb[0:6]}"
                        run_html = f"<span style='color:{color_hex}'>{run_html}</span>"
                    
                    html_para += run_html
                    para_text += run.text
                
                html_para += "</p>"
                html_content += html_para
                plain_text += para_text + "\n"
            
            for table in doc.tables:
                html_content += "<table border='1' cellpadding='3'>"
                for row in table.rows:
                    html_content += "<tr>"
                    for cell in row.cells:
                        html_content += "<td>"
                        for para in cell.paragraphs:
                            cell_text = ""
                            for run in para.runs:
                                text = html.escape(run.text)
                                run_html = text
                                
                                if run.bold:
                                    run_html = f"<strong>{run_html}</strong>"
                                if run.italic:
                                    run_html = f"<em>{run_html}</em>"
                                if run.underline:
                                    run_html = f"<u>{run_html}</u>"
                                if run.font.highlight_color:
                                    run_html = f"<mark>{run_html}</mark>"
                                
                                html_content += run_html
                                cell_text += run.text
                            
                            html_content += "<br/>"
                            plain_text += cell_text + "\n"
                        html_content += "</td>"
                    html_content += "</tr>"
                html_content += "</table>"
                plain_text += "\n"
            
            html_content += "</body></html>"
            
            with open(html_output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            with open(txt_output_path, 'w', encoding='utf-8') as f:
                f.write(plain_text)
            
            print(f"Extraction complete: {html_output_path} and {txt_output_path}")
        except Exception as e:
            print(f"Error extracting text from {doc_path.name}: {e}")
            import traceback
            traceback.print_exc()

if __name__ == "__main__":
    try:
        import docx
    except ImportError:
        print("Installing python-docx...")
        import subprocess
        subprocess.check_call(["pip", "install", "python-docx"])
        import docx
    
    extract_text_with_formatting_from_docx()